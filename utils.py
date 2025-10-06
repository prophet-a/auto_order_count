# utils.py
import docx
import re
import pandas as pd
import os
import json
from datetime import datetime
from openpyxl import load_workbook
from openpyxl.styles import PatternFill
from openpyxl.utils import get_column_letter
# Відкладений імпорт stanza - завантажується тільки коли потрібно
nlp_uk = None

def get_stanza_pipeline():
    """Ледаче завантаження stanza pipeline"""
    global nlp_uk
    if nlp_uk is None:
        try:
            import stanza
            print("Завантаження stanza pipeline...")
            nlp_uk = stanza.Pipeline('uk', processors='tokenize,mwt,pos,lemma', use_gpu=False)
        except Exception as e:
            print(f"Попередження: stanza не завантажено - {e}")
            # Повертаємо None якщо stanza недоступна
            nlp_uk = False
    return nlp_uk if nlp_uk != False else None

def load_doc(file_path):
    """Завантаження документа Word."""
    try:
        return docx.Document(file_path)
    except Exception as e:
        print(f"Помилка завантаження документа {file_path}: {e}")
        return None

def load_ranks(file_path="ranks.txt"):
    """Зчитування звань з файлу та мапування їх на базову форму."""
    rank_map = {}
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            for line in file:
                forms = [r.strip() for r in line.strip().split(",") if r.strip()]
                if forms:
                    base = forms[0]  # нормалізована форма
                    for f in forms:
                        rank_map[f] = base
    except Exception as e:
        print(f"Помилка зчитування з файлу {file_path}: {e}")
    return rank_map

def load_locations(file_path="locations.txt"):
    """Зчитування тригерів для локації."""
    locations = {}
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            for line in file:
                if ':' in line:
                    trigger, loc = line.strip().split(":", 1)
                    locations[trigger.strip().strip('"')] = loc.strip().strip('"')
    except Exception as e:
        print(f"Помилка зчитування локацій з файлу {file_path}: {e}")
    return locations

def parse_date(date_text):
    """
    Розбирає дату з тексту у формат DD.MM.YYYY.
    
    Args:
        date_text (str): Текст з датою у форматі "DD місяць YYYY"
        
    Returns:
        str: Дата у форматі DD.MM.YYYY або None
    """
    if not date_text:
        return None
    
    # Словник відповідності українських місяців до номерів
    months = {
        'січня': '01', 'лютого': '02', 'березня': '03', 'квітня': '04',
        'травня': '05', 'червня': '06', 'липня': '07', 'серпня': '08',
        'вересня': '09', 'жовтня': '10', 'листопада': '11', 'грудня': '12'
    }
    
    # Замінюємо спеціальні символи на пробіли
    date_text = re.sub(r'[«»""„"]', ' ', date_text)
    
    # Шукаємо у форматі "число місяць рік"
    match = re.search(r'(\d{1,2})\s+(\w+)\s+(\d{4})', date_text)
    
    if match:
        day = match.group(1).zfill(2)  # день з провідним нулем
        month_name = match.group(2).lower()
        year = match.group(3)
        
        # Знаходимо номер місяця
        if month_name in months:
            month = months[month_name]
            return f"{day}.{month}.{year}"
    
    return None

def save_to_excel_append(data, output_path="results.xlsx"):
    """Збереження результатів в Excel. Якщо файл існує – додаються нові рядки."""
    columns = ["rank", "name", "VCH", "location", "OS", "date_k", "meal", "cause", "action"]
    new_df = pd.DataFrame(data, columns=columns)
    if os.path.exists(output_path):
        try:
            existing_df = pd.read_excel(output_path)
            final_df = pd.concat([existing_df, new_df], ignore_index=True)
        except Exception as e:
            print(f"Помилка при завантаженні існуючих даних: {e}")
            final_df = new_df
    else:
        final_df = new_df

    try:
        final_df.to_excel(output_path, index=False)
        print(f"Результати успішно збережено в {output_path}")
        
        # Застосовуємо форматування до файлу Excel
        if format_excel_file(output_path):
            print(f"Форматування успішно застосовано до {output_path}")
        else:
            print(f"Не вдалося застосувати форматування до {output_path}")
    except Exception as e:
        print(f"Помилка при збереженні в Excel: {e}")

def determine_paragraph_location(paragraph_norm, location_triggers):
    """
    Визначає локацію для абзацу, надаючи пріоритет "НБ".

    Спочатку шукає патерн типу "N навчального батальйону".
    Якщо не знайдено, використовує extract_location для пошуку за тригерами.

    Args:
        paragraph_norm (str): Нормалізований текст абзацу.
        location_triggers (dict): Словник тригерів локацій.

    Returns:
        str: Знайдена локація ("N НБ" або за тригером) або None.
    """
    # 1. Пріоритетний пошук "N навчального батальйону"
    nb_match = re.search(r'(\d+)[\s-]?(?:го|й|ї)?\s+навчальн(?:ого|ий|ому)\s+батальйон(?:у)?', paragraph_norm, re.IGNORECASE)
    if nb_match:
        location = f"{nb_match.group(1)} НБ"
        print(f"      DEBUG (determine_location): Знайдено пріоритетну локацію НБ: {location}")
        return location

    # 2. Якщо НБ не знайдено, використовуємо загальні тригери
    location = extract_location(paragraph_norm, location_triggers)
    if location:
        print(f"      DEBUG (determine_location): Знайдено локацію за тригером: {location}")
        return location

    # 3. Якщо нічого не знайдено
    print(f"      DEBUG (determine_location): Локація не знайдена ні за НБ, ні за тригерами.")
    return None

def extract_location(text, location_triggers):
    """
    Витягує локацію з тексту на основі тригерних фраз.
    
    Args:
        text (str): Текст для аналізу
        location_triggers (dict): Словник тригерів локацій
        
    Returns:
        str: Знайдена локація або None
    """
    normalized_text = text.lower()
    
    # Перебираємо всі тригери локацій
    for location, triggers in location_triggers.items():
        for trigger in triggers:
            if trigger.lower() in normalized_text:
                print(f"DEBUG: Found location trigger '{trigger}' for location '{location}'")
                return location
    
    return None

def extract_vch(text):
    """
    Витягує номер військової частини з тексту.
    Аліас для функції extract_military_unit з модуля military_personnel.
    
    Args:
        text (str): Текст для аналізу
        
    Returns:
        str: Номер військової частини або None
    """
    from military_personnel import extract_military_unit
    return extract_military_unit(text)

def load_config(config_file='config.json'):
    """
    Завантажує конфігурацію з JSON-файлу.
    
    Args:
        config_file (str, optional): Шлях до файлу конфігурації. За замовчуванням 'config.json'.
        
    Returns:
        dict: Словник з конфігурацією
    """
    try:
        with open(config_file, 'r', encoding='utf-8') as f:
            config = json.load(f)
        return config
    except Exception as e:
        print(f"Помилка при завантаженні конфігурації: {e}")
        # Повертаємо стандартну конфігурацію
        return {
            "rank_map": {},
            "location_triggers": {}
        }

def load_doc(file_path):
    """
    Завантажує документ Word і повертає його текст.
    
    Args:
        file_path (str): Шлях до .docx файлу
        
    Returns:
        str: Текст документа
    """
    try:
        doc = docx.Document(file_path)
        full_text = []
        
        for para in doc.paragraphs:
            full_text.append(para.text)
            
        return '\n'.join(full_text)
    except Exception as e:
        print(f"Помилка при завантаженні документа: {e}")
        return ""

def format_excel_file(excel_file_path):
    """
    Застосовує форматування до Excel файлу:
    1. Автоматично регулює ширину колонок відповідно до вмісту (крім колонки VCH)
    2. Підсвічує клітинки в колонці 'action':
       - 'виключити' = світло-червоний
       - 'зарахувати' = світло-зелений
    
    Args:
        excel_file_path (str): Шлях до Excel файлу для форматування
        
    Returns:
        bool: True при успішному форматуванні, False у випадку помилки
    """
    try:
        # Завантажуємо робочу книгу
        wb = load_workbook(excel_file_path)
        ws = wb.active
        
        # Визначаємо кольори для підсвічування
        light_red = PatternFill(start_color='FFCCCC', end_color='FFCCCC', fill_type='solid')
        light_green = PatternFill(start_color='CCFFCC', end_color='CCFFCC', fill_type='solid')
        
        # Знаходимо індекс колонки 'action' та 'VCH'
        action_col_idx = None
        vch_col_idx = None
        
        # Отримуємо заголовки колонок
        headers = [cell.value for cell in ws[1]]
        
        for idx, header in enumerate(headers, start=1):
            if header == 'action':
                action_col_idx = idx
            elif header == 'VCH':
                vch_col_idx = idx
        
        # Застосовуємо кольори до клітинок в колонці 'action'
        if action_col_idx:
            action_col_letter = get_column_letter(action_col_idx)
            for row in range(2, ws.max_row + 1):  # Starting from row 2 to skip header
                cell = ws[f"{action_col_letter}{row}"]
                if row == 1:  # Пропускаємо комірку заголовку
                    pass
                elif cell.value == 'виключити':
                    cell.fill = light_red
                elif cell.value == 'зарахувати':
                    cell.fill = light_green
        
        # Встановлюємо мінімальний верхній відступ (0.25 дюйма)
        # Автоматично регулюємо ширину колонок, крім VCH
        for idx, column_cells in enumerate(ws.columns, start=1):
            if idx != vch_col_idx:  # Пропускаємо колонку VCH
                length = max(len(str(cell.value)) for cell in column_cells if cell.value)
                # Додаємо трохи додаткового простору
                adjusted_width = length + 2
                ws.column_dimensions[get_column_letter(idx)].width = adjusted_width
        
        # Зберігаємо форматований файл
        wb.save(excel_file_path)
        return True
    except Exception as e:
        print(f"Помилка при форматуванні Excel файлу: {e}")
        return False

def save_results(results, output_file='results.json'):
    """
    Зберігає результати у JSON-файл.
    
    Args:
        results (list): Список результатів
        output_file (str, optional): Шлях до файлу виводу. За замовчуванням 'results.json'.
        
    Returns:
        bool: True якщо збереження успішне, False у випадку помилки
    """
    try:
        with open(output_file, 'w', encoding='utf-8') as f:
            json.dump(results, f, ensure_ascii=False, indent=2)
        return True
    except Exception as e:
        print(f"Помилка при збереженні результатів: {e}")
        return False

# --- Stanza Ukrainian Surname Converter ---
def convert_surnames_to_nominative(text):
    """
    Находит фамилии в родительном падеже и возвращает их в именительном падеже.
    """
    pipeline = get_stanza_pipeline()
    if pipeline is None:
        # Якщо stanza недоступна, повертаємо текст без змін
        return text
    doc = pipeline(text)
    new_tokens = []
    for sent in doc.sentences:
        for word in sent.words:
            # Проверяем, что это имя собственное (PROPN) и родительный падеж (Case=Gen)
            if word.upos == 'PROPN' and word.feats and 'Case=Gen' in word.feats:
                # Используем лемму (обычно это именительный падеж)
                new_tokens.append(word.lemma)
            else:
                new_tokens.append(word.text)
    return ' '.join(new_tokens)


def highlight_names_in_documents(input_dir='in', output_dir='out_highlighted', results_file='results.json'):
    """
    Виділяє прізвища у вихідних документах кольором:
    - Зелений фон: зарахувати
    - Червоний фон: виключити
    
    Args:
        input_dir (str): Папка з оригінальними .docx файлами
        output_dir (str): Папка для збереження виділених файлів
        results_file (str): Файл з результатами обробки
    """
    print("\n=== Виділення прізвищ у документах ===")
    
    # Завантажуємо results.json
    if not os.path.exists(results_file):
        print(f"Помилка: Файл {results_file} не знайдено")
        return False
    
    with open(results_file, 'r', encoding='utf-8') as f:
        results = json.load(f)
    
    # Групуємо по дії
    names_to_add = set()  # Зелений
    names_to_remove = set()  # Червоний
    
    for record in results:
        # Використовуємо оригінальне name (до нормалізації), бо в документі саме так
        name = record.get('name', '')
        action = record.get('action', '')
        
        if name:
            # Витягуємо прізвище (перше слово)
            surname = name.split()[0] if name.split() else ''
            if surname:
                if action == 'зарахувати':
                    names_to_add.add(surname.upper())
                elif action == 'виключити':
                    names_to_remove.add(surname.upper())
    
    print(f"Прізвищ для виділення зеленим: {len(names_to_add)}")
    print(f"Прізвищ для виділення червоним: {len(names_to_remove)}")
    
    # Створюємо вихідну папку
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
        print(f"Створено папку: {output_dir}")
    
    # Обробляємо всі .docx файли
    docx_files = [f for f in os.listdir(input_dir) if f.lower().endswith('.docx') and not f.startswith('~$')]
    
    if not docx_files:
        print(f"Не знайдено .docx файлів у папці {input_dir}")
        return False
    
    from docx.shared import RGBColor
    from docx.enum.text import WD_COLOR_INDEX
    
    for filename in docx_files:
        input_path = os.path.join(input_dir, filename)
        output_path = os.path.join(output_dir, filename)
        
        print(f"\nОбробка: {filename}")
        
        try:
            doc = docx.Document(input_path)
            highlighted_count = {'add': 0, 'remove': 0}
            
            # Обробляємо кожен параграф
            for paragraph in doc.paragraphs:
                paragraph_text = paragraph.text.upper()
                
                # Перевіряємо чи параграф містить прізвища
                found_add_surnames = [s for s in names_to_add if s in paragraph_text]
                found_remove_surnames = [s for s in names_to_remove if s in paragraph_text]
                
                # Якщо знайдено хоча б одне прізвище, виділяємо runs
                if found_remove_surnames or found_add_surnames:
                    # Виділяємо кожен run у цьому параграфі
                    for run in paragraph.runs:
                        run_text = run.text.upper()
                        
                        # Перевіряємо чи цей run містить якесь прізвище
                        if any(surname in run_text for surname in found_remove_surnames):
                            # Червоний фон для виключення
                            run.font.highlight_color = WD_COLOR_INDEX.RED
                            highlighted_count['remove'] += 1
                        elif any(surname in run_text for surname in found_add_surnames):
                            # Зелений фон для зарахування
                            run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
                            highlighted_count['add'] += 1
            
            # Обробляємо таблиці (якщо є)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            paragraph_text = paragraph.text.upper()
                            
                            # Перевіряємо чи параграф містить прізвища
                            found_add_surnames = [s for s in names_to_add if s in paragraph_text]
                            found_remove_surnames = [s for s in names_to_remove if s in paragraph_text]
                            
                            # Якщо знайдено хоча б одне прізвище, виділяємо runs
                            if found_remove_surnames or found_add_surnames:
                                for run in paragraph.runs:
                                    run_text = run.text.upper()
                                    
                                    # Перевіряємо чи цей run містить якесь прізвище
                                    if any(surname in run_text for surname in found_remove_surnames):
                                        run.font.highlight_color = WD_COLOR_INDEX.RED
                                        highlighted_count['remove'] += 1
                                    elif any(surname in run_text for surname in found_add_surnames):
                                        run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
                                        highlighted_count['add'] += 1
            
            # Зберігаємо документ
            doc.save(output_path)
            print(f"  ✅ Збережено: {output_path}")
            print(f"     Зелених: {highlighted_count['add']}, Червоних: {highlighted_count['remove']}")
            
        except Exception as e:
            print(f"  ❌ Помилка при обробці {filename}: {e}")
    
    print(f"\n✅ Виділення завершено! Файли збережено у папці: {output_dir}")
    return True
